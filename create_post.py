"""Generate Hebrew social media post as Word document."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default RTL and font
style = doc.styles['Normal']
style.font.name = 'David'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# RTL helper
def add_rtl_paragraph(text, bold=False, size=None, space_after=None):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'David'
    run.font.size = Pt(size if size else 13)
    if bold:
        run.bold = True
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    # Set RTL on run
    rPr = run._element.get_or_add_rPr()
    rtl = rPr.makeelement(qn('w:rtl'), {})
    rPr.append(rtl)
    return p

# --- Title ---
add_rtl_paragraph('פוסט יום השואה 2026 - פייסבוק ולינקדאין', bold=True, size=16, space_after=12)

# --- POST CONTENT ---

add_rtl_paragraph(
    'תקציר הפרקים הקודמים.',
    bold=True, space_after=2
)

add_rtl_paragraph(
    'לפני שנתיים, ביום השואה, מצאתי באתר יד ושם את העדות של סבא שלום, '
    'שלום דוידוביץ ז"ל, סבה של נעמה אשתי. שלוש שעות של עדות מצולמת, '
    'שלו ושל אחותו הדי. לקחתי את התמלול מיוטיוב, הכנסתי ל-Claude, '
    'וביקשתי סיכום וסיפור ילדים. תוך חצי שעה היה לנו "הסיפור של שלום והדי - '
    'מסע של אומץ ואהבה". שיתפנו עם המשפחה ושיתפתי ברשת, עם הצעה שאחרים יעשו את אותו הדבר.'
)

add_rtl_paragraph(
    'שנה שעברה לקחתי את הסיפור צעד קדימה. ChatGPT ייצר איורים לכל פרק, '
    'הוספתי קריינות של AI, ערכתי הכל בקנבה לסרטון קצר. שוב שיתפתי, שוב הצעתי.'
)

add_rtl_paragraph(
    'השנה שאלתי את עצמי שאלה אחרת: מה אם כל המהלך הזה יכול להיות אוטומטי?',
    space_after=10
)

add_rtl_paragraph(
    'אז בניתי סוכן (agent) שעושה את כל התהליך לבד. נותנים לו קישור ליוטיוב או '
    'טקסט מודבק, והוא מוציא בסוף סרטון מאויר עם קריינות בעברית. '
    'הוא מתמלל, כותב סיפור ילדים מותאם גיל, מצייר איורים בצבעי מים, '
    'מקריא בעברית, ומרכיב סרטון עם מעברים ואפקטים. '
    'כן, ברצינות. הכל אוטומטי.'
)

add_rtl_paragraph(
    'הסרטון המצורף הוא הדגמה ראשונית. הוא לא מושלם. האיורים לא תמיד עקביים, '
    'הקריינות בעברית עדיין מגמגמת פה ושם, ויש עוד מה לשפר. '
    'אבל הוא מראה את הכיוון.'
)

add_rtl_paragraph(
    'פרסמתי את הכלי כקוד פתוח בגיטהאב. מי שיש לו Claude Code '
    'ומפתח API חינמי של Gemini יכול להוריד ולנסות. '
    'ההוראות בקישור:',
    space_after=2
)

# Link paragraph
add_rtl_paragraph(
    'https://github.com/elad-refoua/story-to-video-agent',
    space_after=10
)

add_rtl_paragraph(
    'ויש לי בקשה, או יותר נכון תקווה. '
    'ביד ושם יש אלפי עדויות מצולמות. רובן שוכבות באתר, וצריך לדעת בדיוק מה לחפש כדי למצוא אותן. '
    'הייתי שמח לראות ארגונים כמו יד ושם שוקלים לפתוח ממשקים (API) '
    'שיאפשרו לכלים אוטומטיים לגשת לעדויות בצורה מסודרת, '
    'כך שמשפחות יוכלו ליצור סרטונים כאלה בקלות, בלי לחפש, '
    'בלי להתמודד עם תמלולים אוטומטיים גרועים, בלי שזה ידרוש ידע טכני. '
    'אני לא יודע מה הדרך הנכונה לעשות את זה. אבל הכיוון שווה מחשבה.'
)

add_rtl_paragraph(
    'סבא שלום רצה שיזכרו. שיספרו. שהנכדים ונכדי הנכדים ידעו מאיפה הם באו. '
    'הכלים לעשות את זה קיימים היום. הם לא מושלמים, אבל הם כאן. '
    'והשאלה היא לא "האם הטכנולוגיה מוכנה". השאלה היא האם אנחנו ניקח חצי שעה ונעשה את זה, '
    'בשביל מישהו שכבר לא יכול לספר בעצמו.'
)

# Hashtags - minimal, as requested
add_rtl_paragraph('')
add_rtl_paragraph(
    '#YomHaShoah #HolocaustRemembrance #OpenSource',
    space_after=0
)

# --- MARGIN COMMENT: Editorial notes ---
# Add a comment-like note as a separate section
add_rtl_paragraph('')
add_rtl_paragraph('---', space_after=4)
p_note = add_rtl_paragraph(
    'הערות עריכה (לא לפרסום):',
    bold=True, size=11
)

notes = [
    'ספירת מילים: כ-310 מילים, בתוך מגבלת 400.',
    'הפוסט מתאים לפייסבוק ולינקדאין כמות שהוא. בפייסבוק ניתן להוריד את ההאשטגים.',
    'הסרטון צריך להיות מצורף כקובץ או כקישור בתגובה ראשונה.',
    'מבנה: סיפור אישי > כלי > דוגמה > הזמנה > בקשה > למה.',
    'ביטוי "מגמגמת פה ושם" - ביטוי כנה שמאפשר הערכה ריאלית בלי התנצלות.',
]
for note in notes:
    add_rtl_paragraph(f'- {note}', size=10)

# Save
output_path = r'C:\Users\user\Desktop\projects\movie\story-to-video-agent\post_yom_hashoah_2026.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
